import os
import polars as pl
import pdfplumber
import pandas as pd
from typing import Dict, Any, Callable
from app.tools.base import BaseNode
from app.utils.semantic_profiler import profile_and_cast_df

class FileInputNode(BaseNode):
    """
    FileInputNode ingests data from local files.
    
    COMMUNITY EXTENSIBILITY GUIDE:
    To add support for a new file type (e.g., JSON, Parquet):
    1. Create a new method (e.g., `_parse_json(self, file_path: str) -> pl.DataFrame`).
    2. Register the extension mapping inside `_get_parser_registry()`.
    """
    
    MANIFEST = {
        "id": "fileInput",
        "name": "File Input",
        "category": "inout",
        "icon": "FileText",
        "description": "Read data from local CSV, Excel, PDF, Text, or Word files.",
        "ui_schema": [
            {"field": "filePath", "type": "string", "label": "File Path / Name", "default": ""},
            {"field": "fileType", "type": "select", "label": "File Type", "options": ["auto", "csv", "excel", "parquet", "pdf", "text", "word"], "default": "auto"},
            {"field": "process_local", "type": "boolean", "label": "Process Media Locally (OCR/Parse vs AI Pass-through)", "default": True},
            {"field": "unleash_hardware", "type": "boolean", "label": "🚀 Unleash Maximum Local Memory (Bypass Frontend Safeguards)", "default": False}
        ]
    }

    def execute(self, inputs: Dict[str, pl.DataFrame]) -> pl.DataFrame:
        raw_path = self.parameters.get("filePath", "")
        file_type_setting = self.parameters.get("fileType", "auto").lower()
        process_local = self.parameters.get("process_local", True)
        
        # If the file path is a relative path or just a filename, assume it is in the uploads directory
        if raw_path and not os.path.isabs(raw_path):
            base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "uploads"))
            file_path = os.path.join(base_dir, raw_path)
        else:
            file_path = raw_path

        is_wildcard = "*" in file_path or "?" in file_path

        import glob
        if is_wildcard:
            self.log(f"Wildcard detected. Resolving files for pattern: {file_path}")
            files = glob.glob(file_path)
            if not files:
                raise FileNotFoundError(f"No files matched wildcard: {file_path}")
        else:
            if not file_path or not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            files = [file_path]

        registry = self._get_parser_registry()
        
        dfs = []
        for fp in files:
            self.log(f"Starting file read for path: {fp}")
            
            # Resolve file type mapping for this specific file
            current_file_type = file_type_setting
            if current_file_type == "auto":
                ext = os.path.splitext(fp)[1].lower()
                current_file_type = self._resolve_auto_type(ext, registry)
                self.log(f"Auto-detected file type: {current_file_type}")

            # Override behavior if user wants to pass media to Gemini instead of local parsing
            if not process_local and current_file_type in ["image", "pdf"]:
                self.log("Process Locally is OFF. Routing media file to multimodal pass-through mode.")
                current_file_type = "multimodal"

            if current_file_type not in registry:
                raise ValueError(f"Unsupported file type: {current_file_type}. Supported types: {list(registry.keys())}")

            # Execute parser
            parser_func = registry[current_file_type]
            try:
                df = parser_func(fp)
                # Add a metadata column to track the source file
                df = df.with_columns(pl.lit(os.path.basename(fp)).alias("Source_File"))
                dfs.append(df)
            except Exception as e:
                self.log(f"Failed to read {fp}: {e}")
                raise RuntimeError(f"Failed to read {fp}: {e}")
                
        if len(dfs) == 1:
            combined_df = dfs[0]
        else:
            self.log(f"Concatenating {len(dfs)} files...")
            try:
                combined_df = pl.concat(dfs, how="diagonal")
                self.log(f"Successfully combined {len(dfs)} files. Total rows: {combined_df.height}, Columns: {combined_df.width}")
            except Exception as e:
                self.log(f"Error concatenating files: {e}")
                raise ValueError(f"Could not combine files. They may have incompatible schemas: {e}")
                
        # Apply user-defined schema overrides
        schema_overrides = self.parameters.get("schemaOverrides", {})
        if schema_overrides:
            self.log(f"Applying schema overrides: {schema_overrides}")
            cast_exprs = []
            for col, t in schema_overrides.items():
                if col in combined_df.columns:
                    target_type = str(t).lower()
                    if "int" in target_type:
                        cast_exprs.append(pl.col(col).cast(pl.Int64, strict=False))
                    elif "percent" in target_type:
                        expr = pl.col(col).cast(pl.Utf8)
                        expr = expr.str.replace_all(r'[%]', '').cast(pl.Float64, strict=False) / 100.0
                        cast_exprs.append(expr)
            if cast_exprs:
                combined_df = combined_df.with_columns(cast_exprs)
                
        # Run Semantic Profiler to detect Currency, Percentages, and Accounting formats
        self.log("Running Semantic Data Profiler...")
        final_df, semantic_meta = profile_and_cast_df(combined_df, ignore_cols=list(schema_overrides.keys()) if schema_overrides else None)
        if semantic_meta:
            self.log(f"Detected semantic types: {semantic_meta}")
            
        # Store metadata in self so it can be picked up by the execution engine
        self._semantic_metadata = semantic_meta
        
        # Store the complete, accurate dataset so the background pipeline isn't truncated
        self._full_processed_dataframe = final_df

        # Intelligent Frontend Guard Layer: Apply truncation strictly to the returned UI payload
        unleash_hardware = self.parameters.get("unleash_hardware", False)
        if final_df.height > 10000 and not unleash_hardware:
            self.log("⚠️ FRONTEND SAFEGUARD ACTIVE: Restricting preview streaming payload to top 10000 rows for UI rendering stability. Full dataset remains fully intact in-memory via Polars.")
            return final_df.head(10000)
        
        return final_df

    def get_full_dataframe(self) -> pl.DataFrame:
        """Returns the fully processed background dataframe, untruncated."""
        return getattr(self, "_full_processed_dataframe", pl.DataFrame())

    def _get_parser_registry(self) -> Dict[str, Callable[[str], pl.DataFrame]]:
        """Registry mapping file type identifiers to their parsing strategies."""
        return {
            "csv": self._parse_csv,
            "excel": self._parse_excel,
            "parquet": self._parse_parquet,
            "pdf": self._parse_pdf,
            "text": self._parse_txt,
            "word": self._parse_word
        }

    def _resolve_auto_type(self, ext: str, registry: dict) -> str:
        """Map file extensions to registered parser types."""
        if ext == ".csv": return "csv"
        if ext in [".xls", ".xlsx", ".xlsm", ".xlsb", ".ods"]: return "excel"
        if ext in [".parquet", ".arrow"]: return "parquet"
        if ext == ".pdf": return "pdf"
        if ext == ".txt": return "text"
        if ext in [".doc", ".docx"]: return "word"
        return "csv" # Fallback

    def _parse_csv(self, file_path: str) -> pl.DataFrame:
        delimiter = self.parameters.get("csvDelimiter", ",")
        has_header = self.parameters.get("csvHeader", True)
        encoding = self.parameters.get("csvEncoding", "utf-8")
        
        self.log(f"Parsing CSV with delimiter='{delimiter}', encoding='{encoding}', header={has_header}")
        df = pl.read_csv(
            file_path,
            separator=delimiter,
            has_header=has_header,
            encoding=encoding,
            infer_schema_length=100000,
            null_values=["", "NA", "NaN", "null"],
            try_parse_dates=True
        )
        self.log(f"Successfully read CSV. Row count: {df.height}, Column count: {df.width}")
        return df

    def _parse_excel(self, file_path: str) -> pl.DataFrame:
        sheet_name = self.parameters.get("excelSheet", "")
        self.log(f"Parsing Excel file. Sheet: '{sheet_name if sheet_name else 'First Sheet'}'")
        
        if sheet_name:
            df = pl.read_excel(file_path, sheet_name=sheet_name, engine="calamine")
        else:
            df = pl.read_excel(file_path, engine="calamine")
            
        self.log(f"Successfully read Excel. Row count: {df.height}, Column count: {df.width}")
        return df

    def _parse_parquet(self, file_path: str) -> pl.DataFrame:
        self.log(f"Parsing Parquet file: {file_path}")
        df = pl.read_parquet(file_path)
        self.log(f"Successfully read Parquet. Row count: {df.height}, Column count: {df.width}")
        return df

    def _parse_pdf(self, file_path: str) -> pl.DataFrame:
        extraction_mode = self.parameters.get("pdfExtractionMode", "text").lower()
        self.log(f"Parsing PDF using mode: '{extraction_mode}' via pdfplumber...")
        tables_data = []
        headers = None
        
        if extraction_mode == "tables":
            with pdfplumber.open(file_path) as pdf:
                self.log(f"PDF contains {len(pdf.pages)} pages")
                for i, page in enumerate(pdf.pages):
                    page_tables = page.extract_tables()
                    if page_tables:
                        self.log(f"Page {i+1}: extracted {len(page_tables)} table(s)")
                    
                    for table in page_tables:
                        if not table: continue
                        
                        start_row = 0
                        if not headers:
                            raw_headers = table[0]
                            headers = [str(h).strip() if h and str(h).strip() else f"Column_{idx}" for idx, h in enumerate(raw_headers)]
                            start_row = 1
                            self.log(f"Detected columns from PDF table: {headers}")

                        for row in table[start_row:]:
                            cleaned_row = [str(val).strip() if val is not None else "" for val in row[:len(headers)]]
                            if len(cleaned_row) < len(headers):
                                cleaned_row.extend([""] * (len(headers) - len(cleaned_row)))
                            tables_data.append(cleaned_row)

        if not tables_data:
            if extraction_mode == "tables":
                self.log("No tables found in PDF. Falling back to line text extraction...")
            
            with pdfplumber.open(file_path) as pdf:
                text_lines = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_lines.extend(text.splitlines())
            
            if text_lines:
                headers = ["Text_Line"]
                tables_data = [[line.strip()] for line in text_lines if line.strip()]
                self.log(f"Extracted {len(tables_data)} text lines as a single column.")
            else:
                raise ValueError("No text or tables could be extracted from this PDF.")
        pdf_pd = pd.DataFrame(tables_data, columns=headers)
        df = pl.from_pandas(pdf_pd)
        
        self.log(f"Successfully read PDF. Row count: {df.height}, Column count: {df.width}")
        return df

    def _parse_txt(self, file_path: str) -> pl.DataFrame:
        self.log(f"Parsing Text file: {file_path}")
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            lines = [line.strip() for line in f if line.strip()]
            
        if not lines:
            self.log("Text file is empty.")
            return pl.DataFrame({"Text": []})
            
        df = pl.DataFrame({"Text": lines})
        self.log(f"Successfully read Text. Row count: {df.height}, Column count: {df.width}")
        return df
        
    def _parse_word(self, file_path: str) -> pl.DataFrame:
        self.log(f"Parsing Word Document: {file_path}")
        try:
            from docx import Document
        except ImportError:
            self.log("python-docx is missing. Attempting dynamic install...")
            import subprocess
            import sys
            try:
                subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"], stdout=subprocess.DEVNULL)
                from docx import Document
            except Exception as e:
                raise RuntimeError(f"Could not install python-docx. Please install it manually: {e}")
                
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        tables_data = []
        for table in doc.tables:
            for row in table.rows:
                tables_data.append([cell.text.strip() for cell in row.cells])
                
        # If there are tables, we can try to return the first table or a combined table view
        if tables_data:
            self.log(f"Found {len(doc.tables)} tables in the Word document. Extracting tables...")
            headers = tables_data[0]
            data = tables_data[1:]
            df = pl.DataFrame(data, schema=[f"Col_{i}" if not h else str(h) for i, h in enumerate(headers)], orient="row")
        else:
            self.log("No tables found. Extracting paragraphs instead.")
            df = pl.DataFrame({"Paragraph": paragraphs})
            
        self.log(f"Successfully read Word Doc. Row count: {df.height}, Column count: {df.width}")
        return df

